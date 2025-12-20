from flask import Flask, render_template, request, Response
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.stem import PorterStemmer
from nltk import pos_tag
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from docx import Document
from PyPDF2 import PdfReader
import io

# Downloads
nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')

app = Flask(__name__)
stemmer = PorterStemmer()

LAST_EXPORT = {}

# -------- POS TAG FULL FORMS --------
POS_MAP = {
    "NN": "Noun (Singular)",
    "NNS": "Noun (Plural)",
    "NNP": "Proper Noun (Singular)",
    "NNPS": "Proper Noun (Plural)",
    "VB": "Verb (Base Form)",
    "VBD": "Verb (Past Tense)",
    "VBG": "Verb (Gerund)",
    "VBN": "Verb (Past Participle)",
    "VBP": "Verb (Present)",
    "VBZ": "Verb (3rd Person Singular)",
    "JJ": "Adjective",
    "RB": "Adverb",
    "IN": "Preposition",
    "DT": "Determiner",
    "PRP": "Pronoun"
}

# -------- FILE TEXT EXTRACTION --------
def extract_text(file):
    if file.filename.endswith(".pdf"):
        reader = PdfReader(file)
        return " ".join(page.extract_text() or "" for page in reader.pages)

    if file.filename.endswith(".docx"):
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs)

    return ""


@app.route("/", methods=["GET", "POST"])
def index():
    global LAST_EXPORT
    output = None
    text = ""

    if request.method == "POST":

        uploaded = request.files.get("file")
        if uploaded and uploaded.filename:
            text = extract_text(uploaded)
        else:
            text = request.form.get("text", "").strip()

        operation = request.form.get("operation")

        if not text:
            return render_template("index.html", output=None, text="")

        tokens = word_tokenize(text.lower())

        # -------- TOKENIZATION --------
        if operation == "tokenize":
            output = {
                "operation": "Tokenization",
                "tokens": tokens
            }
            LAST_EXPORT = output

        # -------- STEMMING --------
        elif operation == "stem":
            stems = [stemmer.stem(w) for w in tokens]
            output = {
                "operation": "Stemming",
                "stems": stems
            }
            LAST_EXPORT = output

        # -------- LEMMATIZATION --------
        elif operation == "lemmatize":
            lemmas = [stemmer.stem(w) for w in tokens]
            output = {
                "operation": "Lemmatization",
                "lemmatized": lemmas
            }
            LAST_EXPORT = output

        # -------- POS TAGGING --------
        elif operation == "pos":
            tagged = pos_tag(tokens)
            readable = [(w, POS_MAP.get(t, t)) for w, t in tagged]
            output = {
                "operation": "POS Tagging",
                "pos_tags": readable
            }
            LAST_EXPORT = output

        # -------- BOW / TF-IDF --------
        elif operation in ["bow", "tfidf"]:

            # Sentence-wise documents
            documents = sent_tokenize(text)

            if operation == "bow":
                vec = CountVectorizer()
                title = "Bag of Words"
            else:
                vec = TfidfVectorizer(norm="l2")
                title = "TF-IDF"

            matrix = vec.fit_transform(documents)
            vocab = vec.get_feature_names_out().tolist()
            mat = matrix.toarray().tolist()

            output = {
                "operation": title,
                "vocabulary": vocab,
                "matrix": mat
            }

            LAST_EXPORT = {
                "operation": title,
                "vocab": vocab,
                "matrix": mat
            }

    return render_template("index.html", output=output, text=text)


# -------- WORD EXPORT --------
@app.route("/export")
def export_word():
    if not LAST_EXPORT:
        return "No data to export"

    doc = Document()
    doc.add_heading("NLP Text Analysis Result", level=1)
    doc.add_heading(LAST_EXPORT.get("operation", ""), level=2)

    if "tokens" in LAST_EXPORT:
        doc.add_paragraph(", ".join(LAST_EXPORT["tokens"]))

    elif "stems" in LAST_EXPORT:
        doc.add_paragraph(", ".join(LAST_EXPORT["stems"]))

    elif "lemmatized" in LAST_EXPORT:
        doc.add_paragraph(", ".join(LAST_EXPORT["lemmatized"]))

    elif "pos_tags" in LAST_EXPORT:
        for w, t in LAST_EXPORT["pos_tags"]:
            doc.add_paragraph(f"{w} → {t}")

    elif "vocab" in LAST_EXPORT:
        table = doc.add_table(rows=1, cols=len(LAST_EXPORT["vocab"]) + 1)
        hdr = table.rows[0].cells
        hdr[0].text = "Sentence"

        for i, w in enumerate(LAST_EXPORT["vocab"]):
            hdr[i + 1].text = w

        for i, row in enumerate(LAST_EXPORT["matrix"], start=1):
            cells = table.add_row().cells
            cells[0].text = f"Sentence {i}"
            for j, val in enumerate(row):
                cells[j + 1].text = f"{val:.3f}"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return Response(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=NLP_Result.docx"}
    )


if __name__ == "__main__":
    app.run(debug=True)
